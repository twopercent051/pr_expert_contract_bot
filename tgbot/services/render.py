import os
from datetime import datetime

import docx
from docx.shared import Mm
from num2words import num2words


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs

        for item in inline:
            # print(item.text)
            if key in item.text:
                # print(key)
                item.text = item.text.replace(key, value)


def render_docx(render_dict: dict, cluster_data: list):
    template = render_dict["template"]
    template_document = docx.Document(f"{os.getcwd()}/templates/{template}_ooo.docx")
    variables = render_dict["data"]
    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)
        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)
        section = template_document.sections[0]
        header = section.footer
        header_para = header.paragraphs
        for paragraph in header_para:
            # print(paragraph.text)
            replace_text_in_paragraph(paragraph, variable_key, variable_value)
    if cluster_data:
        template_document = cluster_render(doc=template_document, data=cluster_data)
    template_document.save('result.docx')


def cluster_render(doc, data: list):
    counter = 0
    for paragraph in doc.paragraphs:
        if "THEMES" in paragraph.text:
            replace_text_in_paragraph(paragraph, "THEMES", "")
            break
        counter += 1
    num_data = []
    for i, val in enumerate(data, start=1):
        num_data.append(f"{i}) {val}")
    for i in range(len(num_data), 0, -1):
        new = doc.paragraphs[counter + 1].insert_paragraph_before(num_data[i - 1].strip())
        fmt = new.paragraph_format
        fmt.first_line_indent = Mm(12)
    return doc


def row_value(row: str) -> str:
    if len(row.split("||")) == 2:
        return str(row.split("||")[1].strip())
    else:
        return ""


def data_to_json(template: str, data_str: str) -> dict:
    months = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }
    rows = data_str.split("\n")
    day = str(datetime.today().day) if datetime.today().day > 9 else f"0{datetime.today().day}"
    year = str(datetime.today().year)
    price_int = row_value(rows[8])
    if not price_int.isdigit():
        return None
    price_text = num2words(price_int, lang='ru')
    price = f"{price_int} ({price_text}"
    data_dict = dict(NUM_DEAL=row_value(rows[0]),
                     DAY=day,
                     MONTH=months[datetime.today().month],
                     YEAR=year,
                     COMPANY=row_value(rows[1]),
                     DIR_NAME=row_value(rows[2]),
                     SHORT_NAME=row_value(rows[3]),
                     SITE=row_value(rows[4]),
                     CONTACT=row_value(rows[5]),
                     PHONE=row_value(rows[6]),
                     EMAIL=row_value(rows[7]),
                     PRICE=price,
                     OFF_ADDRESS=row_value(rows[9]),
                     POST_ADDRESS=row_value(rows[10]),
                     INN=row_value(rows[11]),
                     KPP=row_value(rows[12]),
                     BANK_COUNT=row_value(rows[13]),
                     BANK_NAME=row_value(rows[14]),
                     COR_COUNT=row_value(rows[15]),
                     BANK_ID=row_value(rows[16]))
    cluster_data = None
    try:
        if template in ["standard", "optima", "premium"]:
            data_dict["REGION"] = row_value(rows[17])
        if template == "marketing":
            vac_float = round(int(price_int) / 1.2, 2)
            vac_text = num2words(int(vac_float), lang='ru')
            vac_cents_text = f"{int(round(vac_float % 1, 2) * 100)} копеек"
            vac_text = f"{int(vac_float)} ({vac_text}) рублей {vac_cents_text}"
            data_dict["OGRN"] = row_value(rows[17])
            data_dict["OKVED"] = row_value(rows[18])
            data_dict["REGION"] = row_value(rows[19])
            data_dict["VAC"] = vac_text
        if template == "razovaya":
            data_dict["DURATION"] = row_value(rows[17])
            data_dict["REGION"] = row_value(rows[18])
        if template == "klasterizaciya":
            request_cost_int = row_value(rows[18])
            if not request_cost_int.isdigit():
                return None
            request_cost_text = num2words(request_cost_int, lang='ru')
            request_cost = f"{request_cost_int} ({request_cost_text})"
            data_dict["DURATION"] = row_value(rows[17])
            data_dict["REQUEST_COST"] = request_cost
            cluster_data = row_value(rows[19]).split(",")
            # data_dict["THEMES"] = "new row\none more row"
            # cluster_render()
    except IndexError as ex:
        print(ex)
        return None
    render_dict = dict(template=template,
                       data=data_dict)
    today_str = datetime.today().strftime('%d.%m.%Y')
    rus_template = dict(audit="Аудит",
                        optima="Оптима",
                        premium="Премиум",
                        standard="Стандарт",
                        marketing="Ведение рекламы",
                        razovaya="Разовая",
                        klasterizaciya="Кластеризация")
    file_name = f"Договор №{row_value(rows[0])} от {today_str}_ООО {row_value(rows[1])}_{rus_template[template]}"
    render_docx(render_dict=render_dict, cluster_data=cluster_data)
    return file_name
